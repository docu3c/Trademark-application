from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from typing import Dict, Any, List


def export_trademark_opinion_to_word(
    trademark_output: str, web_common_law_output: str = None
) -> str:
    """Export trademark opinion to Word document with proper formatting."""
    print(f"Executing export_trademark_opinion_to_word")
    document = Document()

    # Add main title
    title = document.add_heading("Trademark Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process trademark opinion
    document.add_heading("Trademark Office Opinion", level=1)
    process_opinion_content(document, trademark_output)

    # Conditionally add web common law opinion if provided
    if web_common_law_output:
        document.add_heading("Web Common Law Opinion", level=1)
        process_opinion_content(document, web_common_law_output)

    # Save the document
    filename = (
        "Trademark_Opinion.docx"
        if not web_common_law_output
        else "Combined_Trademark_Opinion.docx"
    )
    document.save(filename)
    return filename


def process_opinion_content(document: Document, content: str) -> None:
    """Helper function to process opinion content with proper markdown conversion."""
    print(f"Executing process_opinion_content")
    lines = content.split("\n")
    current_table = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Handle section headers
        if line.startswith(("Section", "WEB COMMON LAW OPINION")):
            document.add_heading(line, level=2)
            continue

        # Handle tables
        if "|" in line and "---" not in line:
            cells = [cell.strip() for cell in line.split("|") if cell.strip()]

            if current_table is None:
                current_table = document.add_table(rows=1, cols=len(cells))
                current_table.style = "Table Grid"
                hdr_cells = current_table.rows[0].cells
                for i, cell in enumerate(cells):
                    format_cell_text(hdr_cells[i], cell)
            else:
                row_cells = current_table.add_row().cells
                for i, cell in enumerate(cells):
                    format_cell_text(row_cells[i], cell)
        else:
            current_table = None
            p = document.add_paragraph()
            format_paragraph_text(p, line)

            # Enhanced formatting for risk assessment
            if any(keyword in line for keyword in ["Risk Category", "Overall Risk"]):
                for run in p.runs:
                    run.font.size = Pt(12)


def format_cell_text(cell, text: str) -> None:
    """Format text in a table cell with markdown conversion."""
    print(f"Executing format_cell_text")
    paragraph = cell.paragraphs[0]
    format_paragraph_text(paragraph, text)


def format_paragraph_text(paragraph, text: str) -> None:
    """Parse and format paragraph text, handling markdown syntax."""
    print(f"Executing format_paragraph_text")
    # Find all bold text segments (text between double asterisks)
    segments = []
    last_end = 0

    # Use regex to find all bold patterns
    bold_pattern = re.compile(r"\*\*(.*?)\*\*")

    for match in bold_pattern.finditer(text):
        # Add regular text before this bold text
        if match.start() > last_end:
            segments.append((text[last_end : match.start()], False))

        # Add the bold text without asterisks
        segments.append((match.group(1), True))
        last_end = match.end()

    # Add any remaining text
    if last_end < len(text):
        segments.append((text[last_end:], False))

    # Create runs with appropriate formatting
    for segment_text, is_bold in segments:
        if segment_text:
            run = paragraph.add_run(segment_text)
            run.bold = is_bold


def add_conflict_paragraph(document: Document, conflict: Dict[str, Any]) -> None:
    """Add conflict details as paragraphs"""
    print(f"Executing add_conflict_paragraph")
    fields = [
        ("Trademark Name", "trademark_name"),
        ("Trademark Status", "Trademark Status"),
        ("Trademark Owner", "Trademark Owner"),
        ("Trademark Class Number", "Trademark class Number"),
        ("Trademark serial number", "Trademark serial number"),
        ("Trademark registration number", "Trademark registration number"),
        ("Trademark Design phrase", "Trademark design phrase"),
    ]

    for label, key in fields:
        p = document.add_paragraph(f"{label}: {conflict.get(key, 'N/A')}")
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.space_after = Pt(0)

    # Add blank line
    p = document.add_paragraph(" ")
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(0)

    # Add reasoning
    if "reasoning" in conflict:
        p = document.add_paragraph(f"{conflict['reasoning']}\n")
        p.paragraph_format.line_spacing = Pt(18)

    # Add final blank line
    p = document.add_paragraph(" ")
    p.paragraph_format.line_spacing = Pt(18)


def add_conflict_paragraph_to_array(conflict: Dict[str, Any]) -> List[str]:
    """Convert conflict details to array of strings"""
    print(f"Executing add_conflict_paragraph_to_array")
    result = []
    fields = [
        ("Trademark Name", "trademark_name"),
        ("Trademark Status", "Trademark Status"),
        ("Trademark Owner", "Trademark Owner"),
        ("Trademark Class Number", "Trademark class Number"),
        ("Trademark serial number", "Trademark serial number"),
        ("Trademark registration number", "Trademark registration number"),
        ("Trademark Design phrase", "Trademark design phrase"),
    ]

    for label, key in fields:
        result.append(f"{label}: {conflict.get(key, 'N/A')}")

    result.append(" ")  # Blank line for spacing
    if "reasoning" in conflict:
        result.append(f"{conflict['reasoning']}\n")
    result.append(" ")  # Blank line for spacing

    return result
