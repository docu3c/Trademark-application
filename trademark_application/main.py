import os

os.environ["STREAMLIT_SERVER_WATCH_DIRS"] = "false"

import streamlit as st
import time
import base64
from io import BytesIO
from typing import Dict, List, Union, Any
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx import Document
import pandas as pd
from pydantic import ValidationError
import fitz
import re

from models.trademark import TrademarkDetails
from services.extraction_service import (
    extract_proposed_trademark_details,
    extract_proposed_trademark_details2,
    extract_trademark_details_code1,
    extract_trademark_details_code2,
    extract_registration_number,
    extract_serial_number,
    extract_design_phrase,
    extract_ownership,
    extract_international_class_numbers_and_goods_services,
)
from services.analysis_service import run_trademark_analysis
from utils.validation import is_correct_format_code1, is_correct_format_code2
from utils.text_processing import preprocess_text, list_conversion

from services.export_service import (
    export_trademark_opinion_to_word,
    add_conflict_paragraph,
    add_conflict_paragraph_to_array,
)
from utils.text_processing import list_conversion
from services.comparison_service import compare_trademarks, assess_conflict
from services.web_common_law_service import (
    web_law_page,
    extract_web_common_law,
    analyze_web_common_law,
)
from utils.formatting_utils import list_conversion
from services.export_service import (
    export_trademark_opinion_to_word,
)


def parse_trademark_details(
    document_path: str,
) -> List[Dict[str, Union[str, List[int]]]]:
    with fitz.open(document_path) as pdf_document:
        all_extracted_data = []
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            page_text = page.get_text()

            if is_correct_format_code1(page_text):
                preprocessed_chunk = preprocess_text(page_text)
                extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                additional_data = (
                    extract_international_class_numbers_and_goods_services(
                        page_text, page_num, pdf_document
                    )
                )
                registration_number = extract_registration_number(
                    page_text, page_num, pdf_document
                )
                serial_number = extract_serial_number(page_text, page_num, pdf_document)
                design_phrase = extract_design_phrase(page_text, page_num, pdf_document)
                ownership_details = extract_ownership(
                    page_text, page_num, proposed_name, pdf_document
                )

                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    extracted_data.update(additional_data)
                    extracted_data["design_phrase"] = design_phrase
                    extracted_data["owner"] = ownership_details
                    extracted_data["serial_number"] = serial_number
                    extracted_data["registration_number"] = registration_number
                    all_extracted_data.append(extracted_data)

                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_name = data.get("trademark_name", "")
                        if "Global Filings" in trademark_name:
                            trademark_name = trademark_name.split("Global Filings")[
                                0
                            ].strip()
                        if re.match(r"^US-\d+", trademark_name):
                            trademark_name = re.sub(
                                r"^US-\d+\s*", "", trademark_name
                            ).strip()
                        status = data.get("status", "").split(",")[0].strip()
                        serial_number = data.get("serial_number", "")
                        owner = data.get("owner", "")
                        international_class_number = data.get(
                            "international_class_numbers", []
                        )
                        goods_services = data.get("goods_services", "")
                        page_number = data.get("page_number", "")
                        registration_number = data.get(
                            "registration_number",
                            "No registration number presented in document",
                        )
                        design_phrase = data.get(
                            "design_phrase", "No Design phrase presented in document"
                        )

                        trademark_details = TrademarkDetails(
                            trademark_name=trademark_name,
                            owner=owner,
                            status=status,
                            serial_number=serial_number,
                            international_class_number=international_class_number,
                            goods_services=goods_services,
                            page_number=page_number,
                            registration_number=registration_number,
                            design_phrase=design_phrase,
                        )
                        trademark_info = {
                            "trademark_name": trademark_details.trademark_name,
                            "owner": trademark_details.owner,
                            "status": trademark_details.status,
                            "serial_number": trademark_details.serial_number,
                            "international_class_number": trademark_details.international_class_number,
                            "goods_services": trademark_details.goods_services,
                            "page_number": trademark_details.page_number,
                            "registration_number": trademark_details.registration_number,
                            "design_phrase": trademark_details.design_phrase,
                        }
                        print(trademark_info)
                        print(
                            "_____________________________________________________________________________________________________________________________"
                        )
                        trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")

            else:
                if not is_correct_format_code2(page_text):
                    continue

                extracted_data = extract_trademark_details_code2(page_text)
                st.info("Code 2")
                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    all_extracted_data.append(extracted_data)

                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_details = TrademarkDetails(
                            trademark_name=data.get("trademark_name", ""),
                            owner=data.get("owner", ""),
                            status=data.get("status", ""),
                            serial_number=data.get("serial_number", ""),
                            international_class_number=data.get(
                                "international_class_number", []
                            ),
                            goods_services=data.get("goods_services", ""),
                            page_number=data.get("page_number", 0),
                            registration_number=data.get("registration_number", ""),
                            design_phrase=data.get("design_phrase", ""),
                        )
                        if (
                            trademark_details.trademark_name != ""
                            and trademark_details.owner != ""
                            and trademark_details.status != ""
                            and trademark_details.goods_services != ""
                        ):
                            trademark_info = {
                                "trademark_name": trademark_details.trademark_name,
                                "owner": trademark_details.owner,
                                "status": trademark_details.status,
                                "serial_number": trademark_details.serial_number,
                                "international_class_number": trademark_details.international_class_number,
                                "goods_services": trademark_details.goods_services,
                                "page_number": trademark_details.page_number,
                                "registration_number": trademark_details.registration_number,
                                "design_phrase": trademark_details.design_phrase,
                            }

                            trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")

        return trademark_list


# Streamlit App
st.title("Trademark Document Parser Version 7.0")

# File upload
uploaded_files = st.sidebar.file_uploader(
    "Choose PDF files", type="pdf", accept_multiple_files=True
)

if uploaded_files:
    if st.sidebar.button("Check Conflicts", key="check_conflicts"):
        total_files = len(uploaded_files)
        progress_bar = st.progress(0)
        # progress_label.text(f"Progress: 0%")  --- Needed to set

        for i, uploaded_file in enumerate(uploaded_files):
            # Save uploaded file to a temporary file path
            temp_file_path = f"temp_{uploaded_file.name}"
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.read())

            start_time = time.time()

            sp = True
            proposed_trademark_details = extract_proposed_trademark_details(
                temp_file_path
            )

            if proposed_trademark_details:
                proposed_name = proposed_trademark_details.get(
                    "proposed_trademark_name", "N"
                )
                proposed_class = proposed_trademark_details.get(
                    "proposed_nice_classes_number"
                )
                proposed_goods_services = proposed_trademark_details.get(
                    "proposed_goods_services", "N"
                )
                if proposed_goods_services != "N":
                    with st.expander(
                        f"Proposed Trademark Details for {uploaded_file.name}"
                    ):
                        st.write(f"Proposed Trademark name: {proposed_name}")
                        st.write(f"Proposed class-number: {proposed_class}")
                        st.write(
                            f"Proposed Goods & Services: {proposed_goods_services}"
                        )
                    class_list = list_conversion(proposed_class)
                else:
                    st.write(
                        "______________________________________________________________________________________________________________________________"
                    )
                    st.write(
                        f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}"
                    )
                    st.write(
                        "______________________________________________________________________________________________________________________________"
                    )
                    sp = False
            else:

                proposed_trademark_details = extract_proposed_trademark_details2(
                    temp_file_path
                )

                if proposed_trademark_details:
                    proposed_name = proposed_trademark_details.get(
                        "proposed_trademark_name", "N"
                    )
                    proposed_class = proposed_trademark_details.get(
                        "proposed_nice_classes_number"
                    )
                    proposed_goods_services = proposed_trademark_details.get(
                        "proposed_goods_services", "N"
                    )
                    if proposed_goods_services != "N":
                        with st.expander(
                            f"Proposed Trademark Details for {uploaded_file.name}"
                        ):
                            st.write(f"Proposed Trademark name: {proposed_name}")
                            st.write(f"Proposed class-number: {proposed_class}")
                            st.write(
                                f"Proposed Goods & Services: {proposed_goods_services}"
                            )
                        class_list = list_conversion(proposed_class)
                    else:
                        st.write(
                            "______________________________________________________________________________________________________________________________"
                        )
                        st.write(
                            f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}"
                        )
                        st.write(
                            "______________________________________________________________________________________________________________________________"
                        )
                        sp = False
                else:
                    st.error(
                        f"Unable to extract Proposed Trademark Details for {uploaded_file.name}"
                    )
                    sp = False
                    continue

            if sp:
                progress_bar.progress(25)

                existing_trademarks = parse_trademark_details(temp_file_path)
                st.write(len(existing_trademarks))

                # Add collapsible sections for debugging
                with st.expander(
                    "Debug: Existing Trademarks (Same Class)", expanded=False
                ):
                    st.write(
                        "Number of trademarks in same class:", len(existing_trademarks)
                    )
                    st.json(existing_trademarks)

                # PRAVEEN WEB COMMON LAW CODE START'S HERE-------------------------------------------------------------------------------------------------------------------------

                # !!! Function used extract the web common law pages into images
                full_web_common_law = web_law_page(temp_file_path)

                progress_bar.progress(50)
                st.success(
                    f"Existing Trademarks Data Extracted Successfully for {uploaded_file.name}!"
                )

                # !!! Function used extract the web common law details from the images using LLM
                extracted_web_law = extract_web_common_law(
                    full_web_common_law, proposed_name
                )

                # New comprehensive analysis
                analysis_result = analyze_web_common_law(
                    extracted_web_law, proposed_name
                )

                # Display results
                with st.expander("Extracted Web Common Law Data"):
                    st.write(extracted_web_law)

                with st.expander("Trademark Legal Analysis"):
                    st.markdown(analysis_result)  # Using markdown for better formatting

                # extracted_web_law ----- Web common law stored in this variable

                # PRAVEEN WEB COMMON LAW CODE END'S HERE-------------------------------------------------------------------------------------------------------------------------

                nfiltered_list = []
                unsame_class_list = []

                # Iterate over each JSON element in trademark_name_list
                for json_element in existing_trademarks:
                    class_numbers = json_element["international_class_number"]
                    # Check if any of the class numbers are in class_list
                    if any(number in class_list for number in class_numbers):
                        nfiltered_list.append(json_element)
                    else:
                        unsame_class_list.append(json_element)

                existing_trademarks = nfiltered_list
                existing_trademarks_unsame = unsame_class_list
                with st.expander(
                    "Debug: Existing Trademarks (Different Class)", expanded=False
                ):
                    st.write(
                        "Number of trademarks in different class:",
                        len(existing_trademarks_unsame),
                    )
                    st.json(existing_trademarks_unsame)

                high_conflicts = []
                moderate_conflicts = []
                low_conflicts = []
                Name_Matchs = []
                no_conflicts = []

                lt = len(existing_trademarks)
                print(f"Existing trademarks: {existing_trademarks}")
                print(f"Existing trademarks unsame: {existing_trademarks_unsame}")

                for existing_trademark in existing_trademarks:
                    conflict = compare_trademarks(
                        existing_trademark,
                        proposed_name,
                        proposed_class,
                        proposed_goods_services,
                    )
                    if conflict is not None:
                        if conflict["conflict_grade"] == "High":
                            high_conflicts.append(conflict)
                        elif conflict["conflict_grade"] == "Moderate":
                            moderate_conflicts.append(conflict)
                        elif conflict["conflict_grade"] == "Low":
                            low_conflicts.append(conflict)
                        else:
                            no_conflicts.append(conflict)

                for existing_trademarks in existing_trademarks_unsame:
                    if existing_trademarks["international_class_number"] != []:
                        conflict = assess_conflict(
                            existing_trademarks,
                            proposed_name,
                            proposed_class,
                            proposed_goods_services,
                        )

                        if conflict["conflict_grade"] == "Name-Match":
                            Name_Matchs.append(conflict)
                        else:
                            print("Low")

                st.sidebar.write("_________________________________________________")
                st.sidebar.subheader("\n\nConflict Grades : \n")
                st.sidebar.markdown(f"File: {proposed_name}")
                st.sidebar.markdown(
                    f"Total number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}"
                )
                st.sidebar.markdown(f"3 conditions satisfied:  {len(high_conflicts)}")
                st.sidebar.markdown(
                    f"2 conditions satisfied:  {len(moderate_conflicts)}"
                )
                st.sidebar.markdown(f"Name Match's Conflicts: {len(Name_Matchs)}")
                st.sidebar.markdown(f"1 condition satisfied: {len(low_conflicts)}")
                st.sidebar.write("_________________________________________________")

                document = Document()

                # Set page size to landscape
                section = document.sections[0]
                new_width, new_height = section.page_height, section.page_width
                section.page_width = new_width
                section.page_height = new_height

                document.add_heading(
                    f"Trademark Conflict List for {proposed_name} (VERSION - 7.0) :"
                )

                document.add_heading("Dashboard :", level=2)
                # document.add_paragraph(f"\n\nTotal number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}\n- High Conflicts: {len(high_conflicts)}\n- Moderate Conflicts: {len(moderate_conflicts)}\n- Name Match's Conflicts: {len(Name_Matchs)}\n- Low Conflicts: {len(low_conflicts)}\n")

                # Updated Calculate the number of conflicts
                total_conflicts = (
                    len(high_conflicts)
                    + len(moderate_conflicts)
                    + len(Name_Matchs)
                    + len(low_conflicts)
                )

                # Create a table with 5 rows (including the header) and 2 columns
                table = document.add_table(rows=5, cols=2)

                # Set the table style and customize the borders
                table.style = "TableGrid"

                tbl = table._tbl
                tblBorders = OxmlElement("w:tblBorders")

                for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                    border_element = OxmlElement(f"w:{border}")
                    border_element.set(qn("w:val"), "single")
                    border_element.set(
                        qn("w:sz"), "4"
                    )  # This sets the border size; you can adjust it as needed
                    border_element.set(qn("w:space"), "0")
                    border_element.set(qn("w:color"), "000000")
                    tblBorders.append(border_element)

                tbl.append(tblBorders)

                # Fill the first column with labels
                labels = [
                    "Total number of conflicts:",
                    "- 3 conditions satisfied:",
                    "- 2 conditions satisfied:",
                    "- Name Match's Conflicts:",
                    "- 1 condition satisfied:",
                ]

                # Fill the second column with the conflict numbers
                values = [
                    total_conflicts,
                    len(high_conflicts),
                    len(moderate_conflicts),
                    len(Name_Matchs),
                    len(low_conflicts),
                ]

                p = document.add_paragraph(" ")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)

                document.add_heading("Trademark Definitions: ", level=2)
                # p = document.add_paragraph(" ")
                # p.paragraph_format.line_spacing = Pt(18)
                p = document.add_paragraph(
                    "CONDITION 1: MARK: NAME-BASED SIMILARITY (comprised of Exact Match, Semantically Equivalent, Phonetically Equivalent, Primary position match)"
                )
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph("CONDITION 2: CLASS: CLASS OVERLAP")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph(
                    "CONDITION 3: GOODS/SERVICES: OVERLAPPING GOODS/SERVICES & TARGET MARKETS"
                )
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph(
                    "DIRECT HIT: Direct Name hit, regardless of the class"
                )
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                p = document.add_paragraph(" ")
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_after = Pt(0)

                # Populate the table with the labels and values
                for i in range(5):
                    table.cell(i, 0).text = labels[i]
                    table.cell(i, 1).text = str(values[i])

                    # Set the font size to 10 for both cells
                    for cell in table.row_cells(i):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

                if len(high_conflicts) > 0:
                    document.add_heading(
                        "Trademarks with 3 conditions satisfied:", level=2
                    )
                    # Create a pandas DataFrame from the JSON list
                    df_high = pd.DataFrame(high_conflicts)
                    df_high = df_high.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_high = document.add_table(
                        df_high.shape[0] + 1, df_high.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_high.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_high.columns):
                        table_high.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_high.iterrows():
                        for j, value in enumerate(row):
                            cell = table_high.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(moderate_conflicts) > 0:
                    document.add_heading(
                        "Trademarks with 2 conditions satisfied:", level=2
                    )
                    # Create a pandas DataFrame from the JSON list
                    df_moderate = pd.DataFrame(moderate_conflicts)
                    df_moderate = df_moderate.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_moderate = document.add_table(
                        df_moderate.shape[0] + 1, df_moderate.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_moderate.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_moderate.columns):
                        table_moderate.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_moderate.iterrows():
                        for j, value in enumerate(row):
                            cell = table_moderate.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(Name_Matchs) > 0:
                    document.add_heading(
                        "Trademarks with Name Match's Conflicts:", level=2
                    )
                    # Create a pandas DataFrame from the JSON list
                    df_Name_Matchs = pd.DataFrame(Name_Matchs)
                    df_Name_Matchs = df_Name_Matchs.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_Name_Matchs = document.add_table(
                        df_Name_Matchs.shape[0] + 1, df_Name_Matchs.shape[1]
                    )
                    # Set a predefined table style (with borders)
                    table_Name_Matchs.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_Name_Matchs.columns):
                        table_Name_Matchs.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_Name_Matchs.iterrows():
                        for j, value in enumerate(row):
                            cell = table_Name_Matchs.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(low_conflicts) > 0:
                    document.add_heading(
                        "Trademarks with 1 condition satisfied:", level=2
                    )
                    # Create a pandas DataFrame from the JSON list
                    df_low = pd.DataFrame(low_conflicts)
                    df_low = df_low.drop(
                        columns=[
                            "Trademark name",
                            "Trademark class Number",
                            "Trademark registration number",
                            "Trademark serial number",
                            "Trademark design phrase",
                            "conflict_grade",
                            "reasoning",
                        ]
                    )
                    # Create a table in the Word document
                    table_low = document.add_table(df_low.shape[0] + 1, df_low.shape[1])
                    # Set a predefined table style (with borders)
                    table_low.style = (
                        "TableGrid"  # This is a built-in style that includes borders
                    )
                    # Add the column names to the table
                    for i, column_name in enumerate(df_low.columns):
                        table_low.cell(0, i).text = column_name
                    # Add the data to the table
                    for i, row in df_low.iterrows():
                        for j, value in enumerate(row):
                            cell = table_low.cell(i + 1, j)
                            cell.text = str(value)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                if len(high_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 3 conditions satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in high_conflicts:
                        add_conflict_paragraph(document, conflict)

                if len(moderate_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 2 conditions satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in moderate_conflicts:
                        add_conflict_paragraph(document, conflict)

                if len(Name_Matchs) > 0:
                    document.add_heading(
                        "Trademarks with Name Match's Conflicts Reasoning:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in Name_Matchs:
                        add_conflict_paragraph(document, conflict)

                if len(low_conflicts) > 0:
                    document.add_heading(
                        "Explanation: Trademarks with 1 condition satisfied:", level=2
                    )
                    p = document.add_paragraph(" ")
                    p.paragraph_format.line_spacing = Pt(18)
                    for conflict in low_conflicts:
                        add_conflict_paragraph(document, conflict)

                conflicts_array = []

                if len(high_conflicts) > 0:
                    conflicts_array.append(
                        "Explanation: Trademarks with 3 conditions satisfied:"
                    )
                    conflicts_array.append(" ")  # Blank line for spacing
                    for conflict in high_conflicts:
                        conflicts_array.extend(
                            add_conflict_paragraph_to_array(conflict)
                        )

                if len(moderate_conflicts) > 0:
                    conflicts_array.append(
                        "Explanation: Trademarks with 2 conditions satisfied:"
                    )
                    conflicts_array.append(" ")  # Blank line for spacing
                    for conflict in moderate_conflicts:
                        conflicts_array.extend(
                            add_conflict_paragraph_to_array(conflict)
                        )

                if len(Name_Matchs) > 0:
                    conflicts_array.append(
                        "Trademarks with Name Match's Conflicts Reasoning:"
                    )
                    conflicts_array.append(" ")  # Blank line for spacing
                    for conflict in Name_Matchs:
                        conflicts_array.extend(
                            add_conflict_paragraph_to_array(conflict)
                        )

                if len(low_conflicts) > 0:
                    conflicts_array.append(
                        "Explanation: Trademarks with 1 condition satisfied:"
                    )
                    conflicts_array.append(" ")  # Blank line for spacing
                    for conflict in low_conflicts:
                        conflicts_array.extend(
                            add_conflict_paragraph_to_array(conflict)
                        )

                # for i in range(70,96):
                #     progress_bar.progress(i)

                progress_bar.progress(100)

                filename = proposed_name
                doc_stream = BytesIO()
                document.save(doc_stream)
                doc_stream.seek(0)
                download_table = f'<a href="data:application/octet-stream;base64,{base64.b64encode(doc_stream.read()).decode()}" download="{filename + " Trademark Conflict Report"}.docx">Download: {filename}</a>'
                st.sidebar.markdown(download_table, unsafe_allow_html=True)
                st.success(
                    f"{proposed_name} Document conflict report successfully completed!"
                )

                opinion_output = run_trademark_analysis(
                    proposed_name,
                    proposed_class,
                    proposed_goods_services,
                    conflicts_array,
                )
                # Ensure extracted_data is defined by assigning the result of extract_web_common_law
                extracted_data = extract_web_common_law(
                    full_web_common_law, proposed_name
                )
                web_common_law_opinion = analyze_web_common_law(
                    extracted_data, proposed_name
                )
                st.write(
                    "------------------------------------------------------------------------------------------------------------------------------"
                )
                st.write(opinion_output)

                # Export to Word
                filename = export_trademark_opinion_to_word(
                    opinion_output, web_common_law_opinion
                )

                # Download button
                with open(filename, "rb") as file:
                    st.sidebar.download_button(
                        label="Download Trademark Opinion",
                        data=file,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

                end_time = time.time()
                elapsed_time = end_time - start_time
                elapsed_time = elapsed_time // 60
                st.write(f"Time taken: {elapsed_time} mins")

                st.write(
                    "______________________________________________________________________________________________________________________________"
                )

        progress_bar.progress(100)
        st.success("All documents processed successfully!")
