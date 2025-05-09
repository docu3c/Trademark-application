import streamlit as st
import time
import base64
from io import BytesIO
from typing import Dict, List, Union, Any
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx import Document
import pandas as pd
import fitz
import re

from models.trademark import TrademarkDetails
from services.extraction_service import (
    extract_proposed_trademark_details,
    extract_proposed_trademark_details2,
    extract_trademark_details_code1,
    extract_trademark_details_code2,
    extract_registration_number,
    extract_serial_number,
    extract_design_phrase,
    extract_ownership,
    extract_international_class_numbers_and_goods_services,
)
from services.enhanced_comparison_service import compare_trademarks_enhanced
from services.web_common_law_service import (
    web_law_page,
    extract_web_common_law,
    analyze_web_common_law,
)
from services.analysis_service import run_trademark_analysis
from services.export_service import (
    export_trademark_opinion_to_word,
    add_conflict_paragraph,
    add_conflict_paragraph_to_array,
)
from utils.validation import is_correct_format_code1, is_correct_format_code2
from utils.text_processing import preprocess_text, list_conversion


def parse_trademark_details(
    document_path: str,
) -> List[Dict[str, Union[str, List[int]]]]:
    """
    Parse trademark details from PDF document.
    This is kept similar to the original implementation for consistency.
    """
    with fitz.open(document_path) as pdf_document:
        all_extracted_data = []
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            page_text = page.get_text()

            if is_correct_format_code1(page_text):
                preprocessed_chunk = preprocess_text(page_text)
                extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                additional_data = (
                    extract_international_class_numbers_and_goods_services(
                        page_text, page_num, pdf_document
                    )
                )
                registration_number = extract_registration_number(
                    page_text, page_num, pdf_document
                )
                serial_number = extract_serial_number(page_text, page_num, pdf_document)
                design_phrase = extract_design_phrase(page_text, page_num, pdf_document)
                ownership_details = extract_ownership(
                    page_text, page_num, proposed_name, pdf_document
                )

                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    extracted_data.update(additional_data)
                    extracted_data["design_phrase"] = design_phrase
                    extracted_data["owner"] = ownership_details
                    extracted_data["serial_number"] = serial_number
                    extracted_data["registration_number"] = registration_number
                    all_extracted_data.append(extracted_data)

                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_name = data.get("trademark_name", "")
                        if "Global Filings" in trademark_name:
                            trademark_name = trademark_name.split("Global Filings")[
                                0
                            ].strip()
                        if re.match(r"^US-\d+", trademark_name):
                            trademark_name = re.sub(
                                r"^US-\d+\s*", "", trademark_name
                            ).strip()

                        trademark_details = TrademarkDetails(
                            trademark_name=trademark_name,
                            owner=data.get("owner", ""),
                            status=data.get("status", "").split(",")[0].strip(),
                            serial_number=data.get("serial_number", ""),
                            international_class_number=data.get(
                                "international_class_numbers", []
                            ),
                            goods_services=data.get("goods_services", ""),
                            page_number=data.get("page_number", ""),
                            registration_number=data.get(
                                "registration_number",
                                "No registration number presented in document",
                            ),
                            design_phrase=data.get(
                                "design_phrase",
                                "No Design phrase presented in document",
                            ),
                        )

                        trademark_info = {
                            "trademark_name": trademark_details.trademark_name,
                            "owner": trademark_details.owner,
                            "status": trademark_details.status,
                            "serial_number": trademark_details.serial_number,
                            "international_class_number": trademark_details.international_class_number,
                            "goods_services": trademark_details.goods_services,
                            "page_number": trademark_details.page_number,
                            "registration_number": trademark_details.registration_number,
                            "design_phrase": trademark_details.design_phrase,
                        }
                        trademark_list.append(trademark_info)
                    except Exception as e:
                        print(f"Error processing trademark {i}: {e}")

            else:
                if not is_correct_format_code2(page_text):
                    continue

                extracted_data = extract_trademark_details_code2(page_text)
                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    all_extracted_data.append(extracted_data)

                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_details = TrademarkDetails(
                            trademark_name=data.get("trademark_name", ""),
                            owner=data.get("owner", ""),
                            status=data.get("status", ""),
                            serial_number=data.get("serial_number", ""),
                            international_class_number=data.get(
                                "international_class_number", []
                            ),
                            goods_services=data.get("goods_services", ""),
                            page_number=data.get("page_number", 0),
                            registration_number=data.get("registration_number", ""),
                            design_phrase=data.get("design_phrase", ""),
                        )

                        if (
                            trademark_details.trademark_name
                            and trademark_details.owner
                            and trademark_details.status
                            and trademark_details.goods_services
                        ):

                            trademark_info = {
                                "trademark_name": trademark_details.trademark_name,
                                "owner": trademark_details.owner,
                                "status": trademark_details.status,
                                "serial_number": trademark_details.serial_number,
                                "international_class_number": trademark_details.international_class_number,
                                "goods_services": trademark_details.goods_services,
                                "page_number": trademark_details.page_number,
                                "registration_number": trademark_details.registration_number,
                                "design_phrase": trademark_details.design_phrase,
                            }
                            trademark_list.append(trademark_info)
                    except Exception as e:
                        print(f"Error processing trademark {i}: {e}")

        return trademark_list


# Streamlit App
st.title("Enhanced Trademark Document Parser")

# File upload
uploaded_files = st.sidebar.file_uploader(
    "Choose PDF files", type="pdf", accept_multiple_files=True
)

if uploaded_files:
    if st.sidebar.button("Check Conflicts", key="check_conflicts"):
        total_files = len(uploaded_files)
        progress_bar = st.progress(0)

        for i, uploaded_file in enumerate(uploaded_files):
            # Save uploaded file to a temporary file path
            temp_file_path = f"temp_{uploaded_file.name}"
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.read())

            start_time = time.time()

            sp = True
            proposed_trademark_details = extract_proposed_trademark_details(
                temp_file_path
            )

            if proposed_trademark_details:
                proposed_name = proposed_trademark_details.get(
                    "proposed_trademark_name", "N"
                )
                proposed_class = proposed_trademark_details.get(
                    "proposed_nice_classes_number"
                )
                proposed_goods_services = proposed_trademark_details.get(
                    "proposed_goods_services", "N"
                )
                if proposed_goods_services != "N":
                    with st.expander(
                        f"Proposed Trademark Details for {uploaded_file.name}"
                    ):
                        st.write(f"Proposed Trademark name: {proposed_name}")
                        st.write(f"Proposed class-number: {proposed_class}")
                        st.write(
                            f"Proposed Goods & Services: {proposed_goods_services}"
                        )
                    class_list = list_conversion(proposed_class)
                else:
                    st.write(
                        "______________________________________________________________________________________________________________________________"
                    )
                    st.write(
                        f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}"
                    )
                    st.write(
                        "______________________________________________________________________________________________________________________________"
                    )
                    sp = False
            else:
                proposed_trademark_details = extract_proposed_trademark_details2(
                    temp_file_path
                )

                if proposed_trademark_details:
                    proposed_name = proposed_trademark_details.get(
                        "proposed_trademark_name", "N"
                    )
                    proposed_class = proposed_trademark_details.get(
                        "proposed_nice_classes_number"
                    )
                    proposed_goods_services = proposed_trademark_details.get(
                        "proposed_goods_services", "N"
                    )
                    if proposed_goods_services != "N":
                        with st.expander(
                            f"Proposed Trademark Details for {uploaded_file.name}"
                        ):
                            st.write(f"Proposed Trademark name: {proposed_name}")
                            st.write(f"Proposed class-number: {proposed_class}")
                            st.write(
                                f"Proposed Goods & Services: {proposed_goods_services}"
                            )
                        class_list = list_conversion(proposed_class)
                    else:
                        st.write(
                            "______________________________________________________________________________________________________________________________"
                        )
                        st.write(
                            f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}"
                        )
                        st.write(
                            "______________________________________________________________________________________________________________________________"
                        )
                        sp = False
                else:
                    st.error(
                        f"Unable to extract Proposed Trademark Details for {uploaded_file.name}"
                    )
                    sp = False
                    continue

            if sp:
                progress_bar.progress(25)

                existing_trademarks = parse_trademark_details(temp_file_path)
                st.write(f"Found {len(existing_trademarks)} existing trademarks")

                # Web Common Law Analysis
                full_web_common_law = web_law_page(temp_file_path)

                progress_bar.progress(50)
                st.success(
                    f"Existing Trademarks Data Extracted Successfully for {uploaded_file.name}!"
                )

                # Extract and analyze web common law
                extracted_web_law = extract_web_common_law(
                    full_web_common_law, proposed_name
                )
                analysis_result = analyze_web_common_law(
                    extracted_web_law, proposed_name
                )

                # Display web common law results
                with st.expander("Extracted Web Common Law Data"):
                    st.write(extracted_web_law)

                with st.expander("Trademark Legal Analysis"):
                    st.markdown(analysis_result)

                # Filter trademarks by class
                nfiltered_list = []
                unsame_class_list = []

                for json_element in existing_trademarks:
                    class_numbers = json_element["international_class_number"]
                    if any(number in class_list for number in class_numbers):
                        nfiltered_list.append(json_element)
                    else:
                        unsame_class_list.append(json_element)

                existing_trademarks = nfiltered_list
                existing_trademarks_unsame = unsame_class_list

                high_conflicts = []
                moderate_conflicts = []
                low_conflicts = []
                name_matches = []
                no_conflicts = []

                # Process trademarks in the same class
                for existing_trademark in existing_trademarks:
                    comparison_result = compare_trademarks_enhanced(
                        existing_trademark,
                        proposed_name,
                        proposed_class,
                        proposed_goods_services,
                    )

                    # Display prominent element analysis
                    with st.expander(
                        f"Analysis for {existing_trademark['trademark_name']}"
                    ):
                        st.write("### Prominent Element Analysis")
                        st.write(comparison_result["phonetic_analysis"])
                        st.write("### Conflict Analysis")
                        st.write(comparison_result["reasoning"])

                    # Categorize conflicts
                    if comparison_result["conflict_grade"] == "High":
                        high_conflicts.append(comparison_result)
                    elif comparison_result["conflict_grade"] == "Moderate":
                        moderate_conflicts.append(comparison_result)
                    elif comparison_result["conflict_grade"] == "Low":
                        low_conflicts.append(comparison_result)
                    else:
                        no_conflicts.append(comparison_result)

                # Process trademarks in different classes
                for existing_trademark in existing_trademarks_unsame:
                    if existing_trademark["international_class_number"] != []:
                        comparison_result = compare_trademarks_enhanced(
                            existing_trademark,
                            proposed_name,
                            proposed_class,
                            proposed_goods_services,
                        )

                        if comparison_result["conflict_grade"] == "Name-Match":
                            name_matches.append(comparison_result)
                            # Display analysis for name matches
                            with st.expander(
                                f"Name Match Analysis for {existing_trademark['trademark_name']}"
                            ):
                                st.write("### Name Match Analysis")
                                st.write(comparison_result["reasoning"])

                # Display summary
                st.sidebar.write("_________________________________________________")
                st.sidebar.subheader("\n\nConflict Grades : \n")
                st.sidebar.markdown(f"File: {proposed_name}")
                st.sidebar.markdown(
                    f"Total number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(name_matches) + len(low_conflicts)}"
                )
                st.sidebar.markdown(f"3 conditions satisfied:  {len(high_conflicts)}")
                st.sidebar.markdown(
                    f"2 conditions satisfied:  {len(moderate_conflicts)}"
                )
                st.sidebar.markdown(f"Name Match's Conflicts: {len(name_matches)}")
                st.sidebar.markdown(f"1 condition satisfied: {len(low_conflicts)}")
                st.sidebar.write("_________________________________________________")

                # Generate Word document
                document = Document()

                # Set page size to landscape
                section = document.sections[0]
                new_width, new_height = section.page_height, section.page_width
                section.page_width = new_width
                section.page_height = new_height

                document.add_heading(
                    f"Trademark Conflict List for {proposed_name} (Enhanced Version) :"
                )

                document.add_heading("Dashboard :", level=2)

                # Create summary table
                total_conflicts = (
                    len(high_conflicts)
                    + len(moderate_conflicts)
                    + len(name_matches)
                    + len(low_conflicts)
                )

                table = document.add_table(rows=5, cols=2)
                table.style = "TableGrid"

                # Add borders
                tbl = table._tbl
                tblBorders = OxmlElement("w:tblBorders")
                for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                    border_element = OxmlElement(f"w:{border}")
                    border_element.set(qn("w:val"), "single")
                    border_element.set(qn("w:sz"), "4")
                    border_element.set(qn("w:space"), "0")
                    border_element.set(qn("w:color"), "000000")
                    tblBorders.append(border_element)
                tbl.append(tblBorders)

                # Populate table
                labels = [
                    "Total number of conflicts:",
                    "- 3 conditions satisfied:",
                    "- 2 conditions satisfied:",
                    "- Name Match's Conflicts:",
                    "- 1 condition satisfied:",
                ]
                values = [
                    total_conflicts,
                    len(high_conflicts),
                    len(moderate_conflicts),
                    len(name_matches),
                    len(low_conflicts),
                ]

                for i in range(5):
                    table.cell(i, 0).text = labels[i]
                    table.cell(i, 1).text = str(values[i])
                    for cell in table.row_cells(i):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

                # Add trademark definitions
                document.add_heading("Trademark Definitions: ", level=2)
                definitions = [
                    "CONDITION 1: MARK: NAME-BASED SIMILARITY (comprised of Exact Match, Semantically Equivalent, Phonetically Equivalent, Primary position match)",
                    "CONDITION 2: CLASS: CLASS OVERLAP",
                    "CONDITION 3: GOODS/SERVICES: OVERLAPPING GOODS/SERVICES & TARGET MARKETS",
                    "DIRECT HIT: Direct Name hit, regardless of the class",
                ]
                for definition in definitions:
                    p = document.add_paragraph(definition)
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)

                # Add detailed conflict sections
                conflict_sections = [
                    (high_conflicts, "Trademarks with 3 conditions satisfied:"),
                    (moderate_conflicts, "Trademarks with 2 conditions satisfied:"),
                    (name_matches, "Trademarks with Name Match's Conflicts:"),
                    (low_conflicts, "Trademarks with 1 condition satisfied:"),
                ]

                for conflicts, title in conflict_sections:
                    if conflicts:
                        document.add_heading(title, level=2)
                        df = pd.DataFrame(conflicts)
                        table = document.add_table(df.shape[0] + 1, df.shape[1])
                        table.style = "TableGrid"
                        for i, column_name in enumerate(df.columns):
                            table.cell(0, i).text = column_name
                        for i, row in df.iterrows():
                            for j, value in enumerate(row):
                                cell = table.cell(i + 1, j)
                                cell.text = str(value)
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(10)

                # Add explanations
                for conflicts, title in conflict_sections:
                    if conflicts:
                        document.add_heading(f"Explanation: {title}", level=2)
                        p = document.add_paragraph(" ")
                        p.paragraph_format.line_spacing = Pt(18)
                        for conflict in conflicts:
                            add_conflict_paragraph(document, conflict)

                # Prepare conflicts array for analysis
                conflicts_array = []
                for conflicts, title in conflict_sections:
                    if conflicts:
                        conflicts_array.append(f"Explanation: {title}")
                        conflicts_array.append(" ")
                        for conflict in conflicts:
                            conflicts_array.extend(
                                add_conflict_paragraph_to_array(conflict)
                            )

                progress_bar.progress(100)

                # Generate and save document
                filename = proposed_name
                doc_stream = BytesIO()
                document.save(doc_stream)
                doc_stream.seek(0)
                download_table = f'<a href="data:application/octet-stream;base64,{base64.b64encode(doc_stream.read()).decode()}" download="{filename + " Trademark Conflict Report"}.docx">Download: {filename}</a>'
                st.sidebar.markdown(download_table, unsafe_allow_html=True)
                st.success(
                    f"{proposed_name} Document conflict report successfully completed!"
                )

                # Generate trademark opinion
                opinion_output = run_trademark_analysis(
                    proposed_name,
                    proposed_class,
                    proposed_goods_services,
                    conflicts_array,
                )

                # Generate web common law opinion
                web_common_law_opinion = analyze_web_common_law(
                    extracted_web_law, proposed_name
                )

                # Display opinions
                st.write(
                    "------------------------------------------------------------------------------------------------------------------------------"
                )
                st.write(opinion_output)

                # Export to Word
                filename = export_trademark_opinion_to_word(
                    opinion_output, web_common_law_opinion
                )

                # Download button
                with open(filename, "rb") as file:
                    st.sidebar.download_button(
                        label="Download Trademark Opinion",
                        data=file,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

                end_time = time.time()
                elapsed_time = end_time - start_time
                elapsed_time = elapsed_time // 60
                st.write(f"Time taken: {elapsed_time} mins")

                st.write(
                    "______________________________________________________________________________________________________________________________"
                )

        progress_bar.progress(100)
        st.success("All documents processed successfully!")
